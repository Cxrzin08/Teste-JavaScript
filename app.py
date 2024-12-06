import os
from flask import Flask, request, render_template, send_file, url_for
from docx import Document
from docx2pdf import convert as word_to_pdf
import pdfplumber
from PyPDF2 import PdfReader
import fitz
from werkzeug.utils import secure_filename

app = Flask(__name__)

# Configuração de pasta e tamanho máximo do arquivo
UPLOAD_FOLDER = os.path.join(os.getcwd(), "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # Limite de 10 MB

def is_valid_extension(filename, valid_extensions):
    """Verifica se o arquivo possui uma extensão válida."""
    return any(filename.lower().endswith(ext) for ext in valid_extensions)

@app.route("/", methods=["GET"])
def index():
    """Página inicial com formulário de upload."""
    return render_template("index.html", download_link=None)

@app.route("/convert", methods=["POST"])
def convert():
    """Converte arquivos entre PDF e Word com base no tipo selecionado."""
    file = request.files.get("file")
    conversion_type = request.form.get("conversionType")

    if not file or not conversion_type:
        return "Arquivo ou tipo de conversão não selecionado.", 400

    input_filename = secure_filename(file.filename)
    input_path = os.path.join(UPLOAD_FOLDER, input_filename)
    file.save(input_path)

    output_filename = f"converted_{input_filename.rsplit('.', 1)[0]}"
    output_path = None

    try:
        if conversion_type == "pdf-to-word":
            if not is_valid_extension(input_filename, [".pdf"]):
                return "Erro: Para PDF para Word, o arquivo enviado deve ser um PDF.", 400
            output_path = os.path.join(UPLOAD_FOLDER, output_filename + ".docx")
            convert_pdf_to_word(input_path, output_path)
        elif conversion_type == "word-to-pdf":
            if not is_valid_extension(input_filename, [".docx"]):
                return "Erro: Para Word para PDF, o arquivo enviado deve ser um DOCX.", 400
            output_path = os.path.join(UPLOAD_FOLDER, output_filename + ".pdf")
            if not convert_word_to_pdf(input_path, output_path):
                return "Erro ao converter o arquivo Word para PDF.", 500
        else:
            return "Tipo de conversão inválido.", 400
    except Exception as e:
        return f"Erro inesperado: {str(e)}", 500

    download_link = url_for("download_file", filename=os.path.basename(output_path))
    return render_template("index.html", download_link=download_link)

@app.route("/download/<filename>")
def download_file(filename):
    """Baixa o arquivo convertido."""
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.exists(file_path):
        return "Arquivo não encontrado.", 404
    return send_file(file_path, as_attachment=True)

def convert_pdf_to_word(input_path, output_path):
    """Converte PDF para Word usando pdfplumber e PyMuPDF como alternativas."""
    errors = []
    try:
        # Primeira tentativa: pdfplumber
        with pdfplumber.open(input_path) as pdf:
            document = Document()
            for page in pdf.pages:
                text = page.extract_text()
                document.add_paragraph(text if text else "[Página sem texto legível]")
            document.save(output_path)
            return
    except Exception as e:
        errors.append(f"Erro com pdfplumber: {str(e)}")

    try:
        # Segunda tentativa: PyMuPDF
        pdf_document = fitz.open(input_path)
        document = Document()
        for page in pdf_document:
            text = page.get_text("text")
            document.add_paragraph(text if text.strip() else "[Página sem texto legível]")
        document.save(output_path)
        return
    except Exception as e:
        errors.append(f"Erro com PyMuPDF: {str(e)}")

    # Todas as tentativas falharam
    raise Exception(f"Todas as tentativas de conversão falharam: {errors}")

def convert_word_to_pdf(input_path, output_path):
    """Converte Word para PDF usando docx2pdf."""
    try:
        word_to_pdf(input_path, os.path.dirname(output_path))
        generated_pdf = input_path.replace(".docx", ".pdf")
        if os.path.exists(generated_pdf):
            os.rename(generated_pdf, output_path)
            return True
    except Exception as e:
        print(f"Erro ao converter Word para PDF: {str(e)}")
    return False

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)), debug=True)
